"""
LaTeX/PDF to DOCX converter.
"""
import re
import io
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def latex_to_docx(latex_content, title='Resume'):
    """Convert LaTeX content to DOCX by parsing LaTeX structure."""
    try:
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        # Parse LaTeX content and build DOCX
        _parse_and_build(doc, latex_content)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    except Exception as e:
        logger.error(f'DOCX conversion failed: {e}')
        return None


def _parse_and_build(doc, latex):
    """Parse LaTeX content and build Word document."""
    # Remove preamble (everything before \begin{document})
    match = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', latex, re.DOTALL)
    body = match.group(1) if match else latex

    # Remove comments
    body = re.sub(r'%.*?\n', '\n', body)

    lines = body.split('\n')
    in_itemize = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Skip LaTeX commands we can't render
        if line.startswith('\\maketitle') or line.startswith('\\pagestyle'):
            continue

        # Handle centering / name
        if '\\begin{center}' in line or '\\end{center}' in line:
            continue

        # Section headers
        section_match = re.match(r'\\section\*?\{(.+?)\}', line)
        if section_match:
            heading = doc.add_heading(_clean_text(section_match.group(1)), level=1)
            heading.style.font.size = Pt(13)
            heading.style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            continue

        # Subsection headers
        subsection_match = re.match(r'\\subsection\*?\{(.+?)\}', line)
        if subsection_match:
            heading = doc.add_heading(_clean_text(subsection_match.group(1)), level=2)
            heading.style.font.size = Pt(11)
            continue

        # Large/huge text (usually name)
        name_match = re.match(r'\\(?:Huge|huge|LARGE|Large|large)\{?\\textbf\{(.+?)\}\}?', line)
        if name_match:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_clean_text(name_match.group(1)))
            run.bold = True
            run.font.size = Pt(18)
            continue

        # Itemize environment
        if '\\begin{itemize}' in line:
            in_itemize = True
            continue
        if '\\end{itemize}' in line:
            in_itemize = False
            continue

        # List items
        item_match = re.match(r'\\item\s*(.*)', line)
        if item_match:
            text = _clean_text(item_match.group(1))
            if text:
                doc.add_paragraph(text, style='List Bullet')
            continue

        # Regular text
        cleaned = _clean_text(line)
        if cleaned and not cleaned.startswith('\\'):
            doc.add_paragraph(cleaned)


def _clean_text(text):
    """Remove LaTeX formatting commands and return clean text."""
    # Handle textbf
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    # Handle textit
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    # Handle emph
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
    # Handle href
    text = re.sub(r'\\href\{[^}]*\}\{([^}]*)\}', r'\1', text)
    # Handle other commands
    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    # Remove remaining backslash commands
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    # Clean up braces
    text = text.replace('{', '').replace('}', '')
    # Clean up multiple spaces
    text = re.sub(r'\s+', ' ', text).strip()
    # Remove hfill, vspace etc
    text = re.sub(r'\\hfill|\\vspace\{[^}]*\}|\\hrule|\\noindent', '', text)
    return text.strip()
